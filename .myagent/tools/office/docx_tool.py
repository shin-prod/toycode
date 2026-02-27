"""Word (DOCX) 操作ツール。python-docx を使用する。"""

import json
import os

from utils.logger import get_logger

logger = get_logger(__name__)


def docx_read(path: str) -> str:
    """DOCX ファイルの内容をテキストとして抽出する。

    段落と表の内容を含む構造情報も JSON で返す。

    Args:
        path: DOCX ファイルのパス

    Returns:
        ドキュメント構造の JSON 文字列
    """
    try:
        import docx
    except ImportError:
        return "エラー: python-docx がインストールされていません。pip install python-docx"

    resolved = os.path.abspath(path)
    if not os.path.isfile(resolved):
        return f"エラー: ファイルが存在しません: {resolved}"

    logger.debug("docx_read: %s", resolved)
    try:
        doc = docx.Document(resolved)

        # 段落情報
        paragraphs = [
            {"index": i, "text": p.text, "style": p.style.name}
            for i, p in enumerate(doc.paragraphs)
        ]

        # テーブル情報
        tables = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append({"table_index": t_idx, "rows": rows})

        result = {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": "\n".join(p["text"] for p in paragraphs if p["text"]),
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("docx_read エラー: %s", e)
        return f"エラー: DOCX 読み込み失敗: {e}"


def docx_edit(path: str, paragraph_index: int, new_text: str) -> str:
    """DOCX の指定インデックスの段落テキストを編集する。

    既存の段落スタイルを保持したままテキストのみを置き換える。

    Args:
        path: DOCX ファイルのパス
        paragraph_index: 編集する段落のインデックス（0始まり）
        new_text: 新しい段落テキスト

    Returns:
        成功メッセージまたはエラーメッセージ
    """
    try:
        import docx
    except ImportError:
        return "エラー: python-docx がインストールされていません。pip install python-docx"

    resolved = os.path.abspath(path)
    if not os.path.isfile(resolved):
        return f"エラー: ファイルが存在しません: {resolved}"

    logger.debug("docx_edit: %s para=%d", resolved, paragraph_index)
    try:
        doc = docx.Document(resolved)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return (
                f"エラー: 段落インデックスが範囲外です "
                f"(0〜{len(doc.paragraphs) - 1})"
            )
        para = doc.paragraphs[paragraph_index]

        # 既存の run をクリアして新しいテキストを設定
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)

        doc.save(resolved)
        return f"段落 {paragraph_index} を更新しました: {resolved}"
    except Exception as e:
        logger.error("docx_edit エラー: %s", e)
        return f"エラー: DOCX 編集失敗: {e}"
